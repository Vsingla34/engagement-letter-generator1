import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import date
import io, os

# For emailing
import yagmail

# For PDF conversion (Windows/Mac only, needs MS Word/Preview)
try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except:
    DOCX2PDF_AVAILABLE = False

# --------- ICAI Engagement Letter Templates (with official structure & wording) ---------- #

ICAI_TEMPLATES = {
    "Statutory Audit": """\
To,
The Board of Directors of {client_name}
{client_address}

Dear Sirs,

I/We refer to your letter dated {appointment_date} regarding my/our appointment as auditors under Section 139 of the Companies Act, 2013, for the financial year beginning {fy_start} and ending {fy_end}.

I/We confirm my/our acceptance and understanding of the audit engagement as set out in this letter.

**Scope & Objective:**  
The audit will be conducted in accordance with the Standards on Auditing (SAs) issued by ICAI, with the objective to express an opinion on whether the financial statements give a true and fair view as required by the Companies Act, 2013 and relevant rules.

**Management’s Responsibilities:**  
- Preparation and presentation of financial statements in accordance with applicable law and standards.
- Maintenance of proper accounting records.
- Implementation of adequate internal financial controls.
- Providing all records and explanations required for the audit.

**Auditor’s Responsibilities:**  
- Conduct the audit as per SAs and relevant legal requirements.
- Issue a report to the members on the financial statements.
- Report any fraud, as per Section 143(12) of the Companies Act, 2013.

{other_terms}

**Fees & Expenses:** {fees}

Kindly acknowledge your acceptance by signing and returning a copy of this letter.

Yours faithfully,  
{signatory_name}  
{designation}  
{firm_name}  
Firm Registration No.: {firm_reg_no}  
Date: {today}  
Place: {place}
""",

    "Tax Audit": """\
To,
The Board of Directors of {client_name}
{client_address}

Dear Sirs,

We refer to your letter dated {appointment_date} regarding our appointment as tax auditors under Section 44AB of the Income-tax Act, 1961 for the financial year {fy_start} to {fy_end}.

**Scope & Objective:**  
We will conduct the tax audit in accordance with Standards on Auditing issued by ICAI and applicable provisions of the Income-tax Act. Our responsibility is to report as required under Form 3CA/3CB and Form 3CD.

**Management’s Responsibilities:**  
- Maintenance of proper books of account and records.
- Providing all explanations, information, and records necessary for the audit.

**Auditor’s Responsibilities:**  
- Audit and reporting in accordance with Income-tax Act and ICAI SAs.
- Reporting any findings as per requirements.

{other_terms}

**Fees & Expenses:** {fees}

Please acknowledge your acceptance.

Yours faithfully,  
{signatory_name}  
{designation}  
{firm_name}  
FRN: {firm_reg_no}  
Date: {today}  
Place: {place}
""",

    "Internal Audit": """\
To,
The Board of Directors of {client_name}
{client_address}

Dear Sirs,

We refer to your letter dated {appointment_date} regarding our appointment as internal auditors for the financial year {fy_start} to {fy_end}.

**Scope & Objective:**  
We will perform the internal audit in accordance with Standards on Internal Audit issued by ICAI. The audit will focus on evaluating and improving risk management, internal control, and governance processes.

**Management’s Responsibilities:**  
- Ensuring proper accounting systems and controls.
- Providing access to all records, personnel, and explanations required.

**Auditor’s Responsibilities:**  
- Report on internal controls and risk management.
- Submit our findings and recommendations.

{other_terms}

**Fees & Expenses:** {fees}

Kindly acknowledge your acceptance.

Yours truly,  
{signatory_name}  
{designation}  
{firm_name}  
FRN: {firm_reg_no}  
Date: {today}  
Place: {place}
""",

    "GST Audit": """\
To,
The Board of Directors of {client_name}
{client_address}

Dear Sirs,

We acknowledge your letter dated {appointment_date} appointing us as auditors under Section 35(5) of the CGST Act, 2017, for FY {fy_start} to {fy_end}.

**Scope & Objective:**  
To conduct the audit as per Section 35(5) of the CGST Act and relevant ICAI guidance.

**Management’s Responsibilities:**  
- Proper maintenance of books and GST records.
- Compliance with GST laws.
- Providing all documents and explanations required.

**Auditor’s Responsibilities:**  
- Reporting as required by the CGST Act and ICAI.
- Sharing our observations and recommendations.

{other_terms}

**Fees & Expenses:** {fees}

Kindly acknowledge by signing and returning a copy of this letter.

Yours truly,  
{signatory_name}  
{designation}  
{firm_name}  
FRN: {firm_reg_no}  
Date: {today}  
Place: {place}
"""
}

# ------------------- Streamlit App ------------------- #

st.set_page_config(page_title="ICAI Engagement Letter Generator", layout="centered")
st.title("📄 ICAI Engagement Letter Generator (Official Format Only)")
st.caption("Auto-generate official ICAI engagement letters for Audit Assignments. | Powered by Streamlit")

# --- Sidebar Firm/Signatory Info ---
st.sidebar.header("Firm & Signatory Info")
firm_name = st.sidebar.text_input("Firm Name", "Singla Vishal & Co.")
firm_reg_no = st.sidebar.text_input("Firm Registration No.", "000000W")
firm_address = st.sidebar.text_area("Firm Address", "W-101, Mangal Bazar, Laxmi Nagar, Delhi")
signatory_name = st.sidebar.text_input("Partner Name", "CA Vishal Singla")
designation = st.sidebar.text_input("Designation", "Partner")
place = st.sidebar.text_input("Place", "Delhi")
logo_file = st.sidebar.file_uploader("Firm Logo (optional)", type=["png", "jpg", "jpeg"])
today_str = date.today().strftime("%d-%m-%Y")

# --- Single Letter Form ---
with st.form("letter_form"):
    st.header("Single Engagement Letter (ICAI Format)")
    assignment_type = st.selectbox("Assignment Type", list(ICAI_TEMPLATES.keys()))
    client_name = st.text_input("Client Name *", "")
    client_address = st.text_area("Client Address *", "")
    fy_start = st.date_input("FY Start", date(date.today().year-1, 4, 1))
    fy_end = st.date_input("FY End", date(date.today().year, 3, 31))
    appointment_date = st.date_input("Date of Appointment Letter", date.today())
    other_terms = st.text_area("Other Relevant Info", "")
    fees = st.text_input("Fees (as per letter)", "To be mutually agreed")
    submitted = st.form_submit_button("Generate Letter")

def generate_docx(letter_text, firm_logo=None):
    doc = Document()
    if firm_logo:
        doc.add_picture(firm_logo, width=Inches(1.5))
    for para in letter_text.split('\n\n'):
        doc.add_paragraph(para)
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def save_temp_docx(docx_bytes):
    temp_path = "temp_engagement_letter.docx"
    with open(temp_path, "wb") as f:
        f.write(docx_bytes.getbuffer())
    return temp_path

def convert_docx_to_pdf(docx_path):
    pdf_path = docx_path.replace(".docx", ".pdf")
    if DOCX2PDF_AVAILABLE:
        docx2pdf_convert(docx_path, pdf_path)
        return pdf_path
    else:
        return None

if submitted:
    # Prepare letter content
    letter_text = ICAI_TEMPLATES[assignment_type].format(
        client_name=client_name,
        client_address=client_address,
        appointment_date=appointment_date.strftime("%d-%m-%Y"),
        fy_start=fy_start.strftime("%d-%m-%Y"),
        fy_end=fy_end.strftime("%d-%m-%Y"),
        other_terms=other_terms,
        fees=fees,
        signatory_name=signatory_name,
        designation=designation,
        firm_name=firm_name,
        firm_reg_no=firm_reg_no,
        today=today_str,
        place=place,
    )
    st.subheader("Preview")
    st.code(letter_text, language="markdown")
    docx_file = generate_docx(letter_text, logo_file)

    st.markdown("#### Download")
    st.download_button("⬇️ Download as Word (.docx)", docx_file, file_name=f"{assignment_type.replace(' ','_')}_Engagement_Letter_{client_name}.docx")

    # PDF Conversion
    temp_docx_path = save_temp_docx(docx_file)
    if DOCX2PDF_AVAILABLE:
        pdf_path = convert_docx_to_pdf(temp_docx_path)
        if pdf_path and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                st.download_button("⬇️ Download as PDF", f, file_name=f"{assignment_type.replace(' ','_')}_Engagement_Letter_{client_name}.pdf")
            os.remove(pdf_path)
    else:
        st.info("PDF download is available only on Windows/Mac (docx2pdf).")

    # ---- Email Integration ----
    with st.expander("📧 Email This Letter"):
        recipient_email = st.text_input("Recipient Email (client)", key="email_to")
        sender_email = st.text_input("Your Gmail Address", key="email_from")
        sender_password = st.text_input("Gmail App Password", type="password", key="email_pwd")
        email_subject = st.text_input("Email Subject", f"Engagement Letter for {assignment_type}")
        email_body = st.text_area("Email Body", f"Dear {client_name},\n\nPlease find attached your engagement letter.\n\nRegards,\n{firm_name}")
        if st.button("Send Email"):
            if recipient_email and sender_email and sender_password:
                docx_file.seek(0)
                yag = yagmail.SMTP(sender_email, sender_password)
                yag.send(
                    to=recipient_email,
                    subject=email_subject,
                    contents=email_body,
                    attachments={f"{assignment_type.replace(' ','_')}_Engagement_Letter_{client_name}.docx": docx_file.getvalue()}
                )
                st.success(f"E-mail sent to {recipient_email}!")
            else:
                st.error("Please fill all email fields.")

# --- Bulk Generation (Excel/CSV) ---
st.divider()
st.header("Bulk Generation (Excel/CSV Upload)")
st.markdown("""
**Instructions:**  
Upload a file with columns: `assignment_type, client_name, client_address, fy_start, fy_end, appointment_date, other_terms, fees`  
Any missing value will be replaced by defaults.
""")
bulk_file = st.file_uploader("Upload Excel/CSV for Bulk Generation", type=['xlsx', 'csv'])

if bulk_file:
    if bulk_file.name.endswith(".csv"):
        df = pd.read_csv(bulk_file)
    else:
        df = pd.read_excel(bulk_file)
    st.write("Preview (first 5 rows):", df.head())
    if st.button("Generate All Engagement Letters (ZIP)"):
        from zipfile import ZipFile
        result_files = []
        zip_path = "letters_bulk.zip"
        with ZipFile(zip_path, 'w') as zipf:
            for idx, row in df.iterrows():
                assignment_type_row = row.get("assignment_type", "Statutory Audit")
                template = ICAI_TEMPLATES.get(assignment_type_row, ICAI_TEMPLATES["Statutory Audit"])
                letter_text = template.format(
                    client_name=row.get("client_name", ""),
                    client_address=row.get("client_address", ""),
                    appointment_date=str(row.get("appointment_date", today_str)),
                    fy_start=str(row.get("fy_start", f"{date.today().year-1}-04-01")),
                    fy_end=str(row.get("fy_end", f"{date.today().year}-03-31")),
                    other_terms=row.get("other_terms", ""),
                    fees=row.get("fees", "To be mutually agreed"),
                    signatory_name=signatory_name,
                    designation=designation,
                    firm_name=firm_name,
                    firm_reg_no=firm_reg_no,
                    today=today_str,
                    place=place,
                )
                docx_file = generate_docx(letter_text, logo_file)
                fname = f"{assignment_type_row.replace(' ','_')}_Engagement_Letter_{row.get('client_name','client')}.docx"
                zipf.writestr(fname, docx_file.getvalue())
        with open(zip_path, "rb") as f:
            st.download_button("Download All Letters (ZIP)", f, file_name="letters_bulk.zip")
        os.remove(zip_path)

st.divider()
st.caption("© 2025 Singla Vishal & Co. | Official ICAI Formats | Powered by Streamlit")
